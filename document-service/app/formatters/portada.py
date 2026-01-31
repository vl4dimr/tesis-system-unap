from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict
from ..config import get_settings


def generar_portada(doc: Document, datos: Dict) -> List[str]:
    """
    Genera o actualiza la portada del documento con los datos proporcionados.

    Datos esperados:
    - titulo: Título de la tesis
    - autor: Nombre del autor
    - asesor: Nombre del asesor
    - escuela: Escuela profesional
    - facultad: Facultad
    - fecha: Fecha de sustentación
    """
    settings = get_settings()
    cambios = []

    # Buscar si ya existe una portada (primera página con universidad)
    portada_encontrada = False
    for i, para in enumerate(doc.paragraphs[:20]):  # Buscar en los primeros 20 párrafos
        texto = para.text.strip().upper()
        if "UNIVERSIDAD" in texto or "NACIONAL" in texto or "ALTIPLANO" in texto:
            portada_encontrada = True
            break

    if portada_encontrada:
        # Actualizar datos en la portada existente
        if "titulo" in datos:
            for para in doc.paragraphs[:30]:
                # Buscar párrafos que podrían ser el título
                texto = para.text.strip()
                if len(texto) > 20 and len(texto) < 300:
                    # Podría ser el título si está en la primera parte
                    for run in para.runs:
                        if run.font.bold or (run.font.size and run.font.size >= Pt(14)):
                            # Este podría ser el título
                            pass
            cambios.append("Portada existente encontrada")
    else:
        cambios.append("No se encontró portada - se recomienda agregar manualmente")

    # Aplicar formato a elementos de portada si existen
    elementos_portada = [
        "UNIVERSIDAD NACIONAL DEL ALTIPLANO",
        "FACULTAD",
        "ESCUELA PROFESIONAL",
        "TESIS",
        "PRESENTADA POR",
        "PARA OPTAR",
        "PUNO",
    ]

    for para in doc.paragraphs[:40]:
        texto = para.text.strip().upper()
        es_elemento_portada = any(elem in texto for elem in elementos_portada)

        if es_elemento_portada:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)

            for run in para.runs:
                run.font.name = settings.font_name

    return cambios
